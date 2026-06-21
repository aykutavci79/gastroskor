"""GastroSkor Teknik Audit Checklist -> Word (.docx)"""
from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

OUT = Path(__file__).with_name("GastroSkor-Teknik-Audit-Checklist.docx")
AUDIT_SNAPSHOT = "21 Haziran 2026"


def add_item(doc: Document, status: str, title: str, detail: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(f"{status}  {title}")
    run.bold = True
    run.font.size = Pt(11)
    if detail:
        d = doc.add_paragraph(detail)
        d.paragraph_format.left_indent = Inches(0.25)
        d.paragraph_format.space_after = Pt(6)
        for r in d.runs:
            r.font.size = Pt(10)
            r.font.color.rgb = RGBColor(0x33, 0x33, 0x33)


def main() -> None:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    title = doc.add_heading("GastroSkor Teknik Audit Checklist", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"Tarih: {date.today().strftime('%d %B %Y')}\n").bold = True
    meta.add_run(
        f"Son guvenlik snapshot: {AUDIT_SNAPSHOT}\n"
        "Kaynak: resto-pulse-mvp/ monorepo | docs/SECURITY_BACKLOG.md (ertelenen isler)\n"
        "Canli: API https://api.gastroskor.com.tr | Web https://www.gastroskor.com.tr"
    )

    doc.add_paragraph()
    legend = doc.add_paragraph(
        "Durum kodlari:  VAR = mevcut ve calisiyor  |  KISMI = var ama eksik/yetersiz  |  "
        "YOK = hic yok  |  TAMAM = 21.06.2026 guvenlik isi tamamlandi  |  "
        "SISTEM = kullanilan teknoloji"
    )
    legend.runs[0].italic = True

    # --- 1. PROJE YAPISI ---
    doc.add_heading("1. Proje Yapisi ve Teknoloji Stack", level=1)
    s1 = [
        ("VAR", "Frontend framework", "Expo SDK ~54.0.33, React Native 0.81.5, React 19.1.0, Expo Router 6"),
        ("VAR", "Backend framework", "FastAPI + Uvicorn (Python)"),
        ("VAR", "Database", "PostgreSQL (SQLAlchemy ORM + Alembic migration)"),
        ("KISMI", "Railway servisleri", "Kodda 1 API servisi (Dockerfile) + PostgreSQL ayri servis; dashboard dogrulanmadi"),
        ("VAR", "Vercel", "Next.js 15 frontend (web + panel + yasal sayfalar); API Railway'de"),
        ("VAR", "State management", "React Context (session-context, app-badges-context); Redux/Zustand yok"),
        ("VAR", "API iletisimi", "REST (JSON over HTTPS, fetch tabanli lib/api.ts)"),
        ("VAR", "Authentication", "Google OAuth (mobil) -> backend JWT (HS256, python-jose); sifre login yok"),
        ("VAR", "Dosya depolama", "Local volume (Railway) veya S3/R2 uyumlu (boto3); menu/yorum/avatar"),
        ("VAR", "Push notification", "expo-notifications + Expo Push Token -> backend /me/push-token"),
    ]
    for st, t, d in s1:
        add_item(doc, st, t, d)

    # --- 2. GUVENLIK ---
    doc.add_heading("2. Guvenlik", level=1)
    doc.add_heading("API Guvenligi", level=2)
    s2a = [
        ("KISMI", "Tum endpoint'ler auth gerektiriyor mu?", "SecurityMiddleware korumali; kasitli public GET endpoint'ler var (restoran listesi, canli yerler)"),
        ("KISMI", "Token olmadan erisilebilir endpoint'ler", "GET /health, /restaurants*, /live/places/*, /regional-flavors/*, /voice-products/catalog, POST /auth/google/*, POST /auth/refresh, POST /panel/applications, GET /media/* — POST /voice/transcribe artik JWT zorunlu"),
        ("VAR", "Rate limiting", "Redis (primary) + in-memory fallback; path kurallari + kullanici global cap; /health rate_limit.redis_ok ile dogrulanir"),
        ("VAR", "CORS", "Whitelist: localhost, gastroskor.vercel.app, gastroskor.com.tr, www.gastroskor.com.tr"),
        ("VAR", "HTTPS", "Railway + Vercel production HTTPS; HTTP redirect Vercel otomatik"),
    ]
    for st, t, d in s2a:
        add_item(doc, st, t, d)

    doc.add_heading("Veri Guvenligi", level=2)
    s2b = [
        ("YOK", "Sifre hash", "Sifre login yok (Google OAuth); passlib requirements'ta ama kullanilmiyor. OTP SHA256 hash"),
        ("KISMI", "Hardcoded secret", "jwt_secret default 'change-me'; production'da RuntimeError ile engelleniyor"),
        ("VAR", ".env gitignore", "mobile/.gitignore ve backend/.env.example mevcut"),
        ("VAR", "Input validation", "Pydantic schemas (backend), TypeScript tipleri (mobil/frontend)"),
        ("VAR", "SQL injection korumasi", "SQLAlchemy ORM parametreli sorgular"),
        ("KISMI", "XSS korumasi", "Next.js otomatik escape; API JSON. Ayri CSP/security headers audit yapilmamis"),
    ]
    for st, t, d in s2b:
        add_item(doc, st, t, d)

    doc.add_heading("Mobil Guvenlik", level=2)
    s2c = [
        ("VAR", "Hassas veri depolama", "JWT + refresh token expo-secure-store (Keychain/Keystore); session ozeti session-secure-storage"),
        ("VAR", "JWT saklama yeri", "SecureStore (auth-token.ts) + bellek cache; AsyncStorage'da token yok"),
        ("VAR", "Token expire", "Access ~2 saat (jwt_access_token_expire_hours: 2); refresh ~30 gun"),
        ("VAR", "Refresh token", "POST /auth/refresh; rotation + revoked_refresh_tokens tablosu (jti iptal listesi)"),
        ("VAR", "Mobil token garantisi", "lib/api.ts request() oncesi ensureAccessToken(); voice-whisper-transcribe ayri cagiriyor"),
    ]
    for st, t, d in s2c:
        add_item(doc, st, t, d)

    doc.add_heading("Tamamlanan Guvenlik Islemleri (21 Haziran 2026)", level=2)
    s2d = [
        ("TAMAM", "Review / feedback IDOR", "resolve_authenticated_email ile JWT != claimed_email engeli; test_author_identity_idor.py"),
        ("TAMAM", "Public voice/transcribe auth acigi", "POST /api/v1/voice/transcribe JWT zorunlu (middleware + require_request_auth); test_voice_transcription.py"),
        ("TAMAM", "Mobil token garantisi", "ensureAccessToken() tum API isteklerinde; Whisper upload oncesi token kontrolu"),
        ("TAMAM", "Refresh token rotation + revocation", "revoked_refresh_tokens migration; POST /auth/refresh yeni cift dondurur; eski jti iptal"),
        ("TAMAM", "SecureStore migrasyonu", "access + refresh token AsyncStorage'dan SecureStore'a tasindi"),
        ("TAMAM", "Redis rate limiting", "REDIS_URL ile dagitik limiter; Railway'de gastroskor:rl* anahtarlari"),
        ("TAMAM", "CI/CD pipeline", "GitHub Actions: backend-ci.yml (pytest + pip-audit), mobile-ci.yml (npm audit high + tsc + test)"),
        ("TAMAM", "npm audit high", "undici + tar duzeltildi; 0 high, 23 moderate (Expo ekosistemi — SECURITY_BACKLOG.md)"),
        ("TAMAM", "Backend test suite", "260/260 pytest (SQLite JSONB/UUID compat conftest dahil)"),
        ("TAMAM", "Mobile TypeScript CI", "tsc --noEmit 0 hata (21.06.2026)"),
    ]
    for st, t, d in s2d:
        add_item(doc, st, t, d)

    # --- 3. VERITABANI ---
    doc.add_heading("3. Veritabani", level=1)
    s3 = [
        ("KISMI", "Disaridan direkt DB baglantisi", "Railway PostgreSQL; public erisim dashboard ayarina bagli"),
        ("KISMI", "DB kullanici yetkisi minimum mu", "Tek app kullanicisi tipik; ayri read-only role yok"),
        ("YOK", "Otomatik yedek", "Railway backup durumu dogrulanamadi; manuel kontrol gerekli"),
        ("VAR", "Index'ler", "Alembic'te users.email, restaurants.city, public_reviews.* vb."),
        ("KISMI", "N+1 query", "Profil edilmemis; buyuk route dosyalarinda risk"),
        ("VAR", "Migration sistemi", "Alembic (39+ migration dosyasi)"),
    ]
    for st, t, d in s3:
        add_item(doc, st, t, d)

    # --- 4. HATA YONETIMI ---
    doc.add_heading("4. Hata Yonetimi", level=1)
    s4 = [
        ("YOK", "Backend global error handler", "FastAPI default + route-level HTTPException; merkezi exception_handler yok"),
        ("VAR", "Anlamli hata mesajlari", "Turkce detail mesajlari; mobilde format-api-error.ts"),
        ("KISMI", "API hata loglama", "Python logging modulu; merkezi log aggregator yok"),
        ("VAR", "Uygulama crash", "AppErrorBoundary (root) + SpeechMicErrorBoundary; crash ekrani"),
        ("KISMI", "Network offline", "format-api-error.ts network mesaji; global offline modu/NetInfo yok"),
        ("VAR", "Crash reporting", "@sentry/react-native prod'da aktif; sentry-scrub PII redaction + beforeSend"),
    ]
    for st, t, d in s4:
        add_item(doc, st, t, d)

    # --- 5. PERFORMANS ---
    doc.add_heading("5. Performans", level=1)
    s5 = [
        ("YOK", "API response suresi", "APM/monitoring kurulu degil"),
        ("KISMI", "Resim optimizasyonu", "expo-image (mobil), Pillow (backend); otomatik resize/CDN sinirli"),
        ("KISMI", "Pagination", "API'de limit parametreleri var; mobilde infinite scroll yok"),
        ("YOK", "API cache", "Redis/in-memory HTTP cache yok"),
        ("KISMI", "Railway instance / auto-scale", "Tek container deploy; auto-scale dogrulanmadi"),
        ("YOK", "En yavas endpoint", "Profil yok; adaylar: AI analiz, Google Places proxy"),
    ]
    for st, t, d in s5:
        add_item(doc, st, t, d)

    # --- 6. KOD KALITESI ---
    doc.add_heading("6. Kod Kalitesi", level=1)
    s6 = [
        ("KISMI", "DRY ihlalleri", "routes.py (1980), panel_routes.py (1286), api.ts (855) satir"),
        ("KISMI", "Tek sorumluluk", "OnlineOrderSection.tsx (634), siparis-acik.tsx (616) fazla sorumluluk"),
        ("KISMI", "Magic number", "Rate limit, STT restart 300ms, limit=20/50 sabitleri daginik"),
        ("VAR", "Console.log temizligi", "console.log bulunamadi; sadece error boundary'lerde error/warn"),
        ("VAR", "TODO / FIXME", "Kodda 0 adet bulundu"),
        ("SISTEM", "En buyuk dosya", "backend/app/api/v1/routes.py - yaklasik 1980 satir"),
        ("KISMI", "Kullanilmayan import", "ESLint var; otomatik temizlik CI'da calismiyor"),
    ]
    for st, t, d in s6:
        add_item(doc, st, t, d)

    # --- 7. TEST ---
    doc.add_heading("7. Test", level=1)
    s7 = [
        ("VAR", "Otomatik test", "Backend: 26+ test dosyasi, 260 pytest (CI'da). Mobil: sentry-scrub.test.ts (node:test)"),
        ("KISMI", "Kritik fonksiyon testleri", "Auth, siparis, OTP, voice catalog testli; odeme entegrasyonu yok"),
        ("KISMI", "Manuel test dokumantasyonu", "STORE.md, ROADMAP.md var; formal test case dokumani yok"),
    ]
    for st, t, d in s7:
        add_item(doc, st, t, d)

    # --- 8. DEPLOYMENT ---
    doc.add_heading("8. Deployment ve DevOps", level=1)
    s8 = [
        ("VAR", "Prod / dev ayri", "environment config; production'da Swagger kapali, dev seed 404"),
        ("VAR", "CI/CD pipeline", ".github/workflows: backend-ci.yml, mobile-ci.yml, sofra-bulmaca-daily.yml"),
        ("KISMI", "Env variable'lar", ".env.example kapsamli; Railway/Vercel dashboard dogrulanmadi"),
        ("YOK", "Railway memory/CPU", "Dashboard bilgisi kodda yok"),
        ("YOK", "Uptime monitoring", "UptimeRobot vb. yok; sadece /health endpoint"),
        ("VAR", "Railway restart policy", "ON_FAILURE, maxRetries: 10 (railway.toml)"),
    ]
    for st, t, d in s8:
        add_item(doc, st, t, d)

    # --- 9. MOBIL ---
    doc.add_heading("9. Mobil Ozel", level=1)
    s9 = [
        ("KISMI", "iOS / Android ozellik paritesi", "Cogu ortak; STT platform davranislari farkli"),
        ("VAR", "Expo SDK guncel", "SDK 54 (current stable)"),
        ("VAR", "Deep link", "gastroskor:// + applinks:www.gastroskor.com.tr"),
        ("KISMI", "Push test (iOS/Android)", "Kod ve entitlement hazir; sistematik test dokumani yok"),
        ("KISMI", "Store metadata", "STORE.md checklist var; tam doldurma belirsiz"),
        ("VAR", "Privacy policy URL", "https://www.gastroskor.com.tr/gizlilik"),
        ("VAR", "Konum izni aciklamasi", "Yakin restoranlari gostermek icin konum kullanilir."),
        ("VAR", "Mikrofon izni aciklamasi", "Gastro Siparis komutlarini konusarak yazmak icin mikrofon kullanilir."),
    ]
    for st, t, d in s9:
        add_item(doc, st, t, d)

    # --- 10. GENEL ---
    doc.add_heading("10. Genel", level=1)
    s10 = [
        ("VAR", "README", "resto-pulse-mvp/README.md + mobile/README.md"),
        ("KISMI", "npm audit", "Mobil: 23 moderate, 0 high (21.06.2026); moderate Expo transitive — SECURITY_BACKLOG.md"),
        ("VAR", "Lisans sorunu", "Bilinen copyleft ihlali yok; formal audit yapilmamis"),
        ("KISMI", "KVKK / GDPR", "Gizlilik/KVKK sayfalari canli; hesap-sil e-posta ile talep; otomatik self-service silme yok"),
    ]
    for st, t, d in s10:
        add_item(doc, st, t, d)

    # --- OZET TABLO ---
    doc.add_page_break()
    doc.add_heading("Ozet Tablo", level=1)
    table = doc.add_table(rows=8, cols=5)
    table.style = "Table Grid"
    headers = ["Kategori", "Toplam Madde", "VAR", "KISMI", "YOK"]
    rows = [
        ["Guvenlik", "20", "12", "5", "0"],
        ["Veritabani", "6", "2", "2", "2"],
        ["Hata Yonetimi", "6", "3", "2", "1"],
        ["Performans", "6", "0", "3", "3"],
        ["Kod Kalitesi", "7", "2", "5", "0"],
        ["Mobil", "9", "5", "4", "0"],
        ["TOPLAM", "54", "24", "21", "6"],
    ]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
    for ri, row in enumerate(rows, start=1):
        for ci, val in enumerate(row):
            table.rows[ri].cells[ci].text = val

    # --- ONCELIKLI AKSIYONLAR ---
    doc.add_heading("Oncelikli Aksiyonlar", level=1)
    actions = [
        ("Yuksek", "DB yedek / PITR yok", "Ilk gercek restoran + online siparis -> Railway Pro + PITR (SECURITY_BACKLOG.md)"),
        ("Orta", "npm audit 23 moderate", "Expo SDK major upgrade ile birlikte npm audit fix --force (SECURITY_BACKLOG.md)"),
        ("Orta", "Global error handler", "FastAPI merkezi exception handler"),
        ("Orta", "API cache / pagination", "Liste endpoint'lerinde cursor pagination"),
        ("Dusuk", "XSS / CSP audit", "Next.js security headers ve CSP politikasi gozden gecir"),
        ("Dusuk", "Uptime monitoring", "UptimeRobot veya benzeri /health disi alarm"),
    ]
    for pri, konu, oneri in actions:
        add_item(doc, pri.upper(), konu, oneri)

    doc.save(OUT)
    print(f"Olusturuldu: {OUT}")


if __name__ == "__main__":
    main()
